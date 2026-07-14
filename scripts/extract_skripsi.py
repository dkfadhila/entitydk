"""Extract ordered thesis content with heading + list hierarchy."""
from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.parts.image import ImagePart

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "scripts" / "skripsi.docx"
OUT_DIR = ROOT / "public" / "evidence" / "images"
OUT_JSON = ROOT / "src" / "lib" / "evidence-content.json"


def save_image(part: ImagePart, saved: dict, counter: list[int]):
    blob = part.blob
    h = hashlib.md5(blob).hexdigest()[:10]
    if h in saved:
        return saved[h]
    counter[0] += 1
    ext = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/jpg": ".jpg",
        "image/gif": ".gif",
        "image/webp": ".webp",
        "image/bmp": ".bmp",
        "image/tiff": ".tiff",
    }.get(part.content_type, ".bin")
    name = f"fig-{counter[0]:03d}-{h}{ext}"
    (OUT_DIR / name).write_bytes(blob)
    info = {
        "src": f"/evidence/images/{name}",
        "name": name,
        "mime": part.content_type,
        "bytes": len(blob),
    }
    saved[h] = info
    return info


def para_images(paragraph: Paragraph, img_map: dict, saved: dict, counter: list[int]):
    images = []
    for blip in paragraph._element.findall(".//" + qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if rid and rid in img_map:
            images.append(save_image(img_map[rid], saved, counter))
    return images


def style_name(paragraph: Paragraph) -> str:
    try:
        return paragraph.style.name if paragraph.style else "Normal"
    except Exception:
        return "Normal"


def list_info(paragraph: Paragraph):
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        return None, None, None
    numPr = pPr.find(qn("w:numPr"))
    ilvl = numId = None
    if numPr is not None:
        il = numPr.find(qn("w:ilvl"))
        ni = numPr.find(qn("w:numId"))
        if il is not None:
            ilvl = int(il.get(qn("w:val")))
        if ni is not None:
            numId = ni.get(qn("w:val"))
    ind = pPr.find(qn("w:ind"))
    left = None
    if ind is not None:
        left = ind.get(qn("w:left"))
        if left is not None:
            left = int(left)
    return ilvl, numId, left


def classify_heading(text: str, style: str):
    t = text.strip()
    s = (style or "").lower()
    if not t:
        return None
    if s.startswith("toc"):
        return None  # skip TOC styles in raw pass; filter later
    if "heading 1" in s or s == "heading 1":
        return "h1"
    if "heading 2" in s or s in ("heading 2", "head2"):
        # head2 often chapter subtitle like PENDAHULUAN
        if s == "head2" or t.isupper():
            return "h1sub"
        return "h2"
    if "heading 3" in s:
        return "h3"
    if "title" in s:
        return "h1"
    if re.match(r"^BAB\s+[IVXLC0-9]+", t, re.I):
        return "h1"
    upper = t.upper()
    markers = [
        "ABSTRAK",
        "ABSTRACT",
        "MOTTO",
        "DAFTAR PUSTAKA",
        "LAMPIRAN",
        "KESIMPULAN",
        "PENDAHULUAN",
        "KAJIAN PUSTAKA",
        "METODE PENELITIAN",
        "HASIL DAN PEMBAHASAN",
    ]
    if len(t) < 90 and any(upper == m or upper.startswith(m) for m in markers):
        if upper in ("PENDAHULUAN", "KAJIAN PUSTAKA", "METODE PENELITIAN", "HASIL DAN PEMBAHASAN", "KESIMPULAN"):
            return "h1sub"
        return "h1"
    return None


def main():
    if not SRC.exists():
        raise SystemExit(f"missing {SRC}")
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for p in OUT_DIR.glob("fig-*"):
        p.unlink()

    doc = Document(str(SRC))
    img_map = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img_map[rel.rId] = rel.target_part

    saved = {}
    counter = [0]
    raw = []

    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            text = p.text.replace("\u00a0", " ").strip()
            images = para_images(p, img_map, saved, counter)
            style = style_name(p)
            ilvl, numId, left = list_info(p)

            if not text and not images:
                continue

            if images and not text:
                for im in images:
                    raw.append({"type": "image", **im})
                continue

            # skip pure TOC styles
            sl = style.lower()
            if sl.startswith("toc") or sl == "table of figures":
                continue

            level = classify_heading(text, style)
            if level == "h1":
                raw.append({"type": "h1", "text": text, "style": style})
            elif level == "h1sub":
                raw.append({"type": "h1sub", "text": text, "style": style})
            elif level == "h2":
                raw.append({"type": "h2", "text": text, "style": style})
            elif level == "h3":
                raw.append({"type": "h3", "text": text, "style": style})
            elif re.match(r"^(Gambar|Figure|Tabel|Table)\s+\d*", text, re.I):
                raw.append({"type": "caption", "text": text, "style": style})
            elif ilvl is not None or (left is not None and left >= 250 and numId is not None):
                # list item — depth from ilvl primarily
                depth = ilvl if ilvl is not None else 0
                # map numId clusters roughly: larger left => deeper
                if ilvl is None and left is not None:
                    depth = 0 if left < 350 else 1 if left < 500 else 2
                raw.append(
                    {
                        "type": "li",
                        "text": text,
                        "depth": int(depth),
                        "numId": numId,
                        "style": style,
                    }
                )
            else:
                raw.append({"type": "p", "text": text, "style": style})

            for im in images:
                raw.append({"type": "image", **im})

        elif child.tag == qn("w:tbl"):
            table = Table(child, doc)
            rows = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = "\n".join(
                        pp.text.strip() for pp in cell.paragraphs if pp.text.strip()
                    )
                    for pp in cell.paragraphs:
                        for im in para_images(pp, img_map, saved, counter):
                            raw.append({"type": "image", **im})
                    cells.append(cell_text)
                rows.append(cells)
            if any(any((c or "").strip() for c in r) for r in rows):
                raw.append({"type": "table", "rows": rows})

    # ---------- FILTER front matter, keep title + abstrak + motto + body ----------
    TITLE = "ANALISIS DINAMIKA ATMOSFER KEJADIAN BANJIR PADA TANGGAL 25 MARET 2025 DI KABUPATEN KUDUS"

    def tx(b):
        return (b.get("text") or "").replace("\t", " ").strip()

    def up(b):
        return tx(b).upper()

    # find markers in raw (after toc skip, indices differ)
    def find_h1(pred):
        for i, b in enumerate(raw):
            if b.get("type") in ("h1", "h1sub") and pred(up(b)):
                return i
        return None

    abstrak_i = find_h1(lambda u: u.startswith("ABSTRAK") and "ii" not in u.lower())
    abstract_i = find_h1(lambda u: u.startswith("ABSTRACT"))
    motto_i = find_h1(lambda u: u == "MOTTO")
    persembahan_i = find_h1(lambda u: "PERSEMBAHAN" in u)
    # first real BAB I body: h1 == BAB I (not toc)
    bab1_i = None
    for i, b in enumerate(raw):
        if b.get("type") == "h1" and re.fullmatch(r"BAB I", tx(b), re.I):
            bab1_i = i
            break
    if bab1_i is None:
        bab1_i = find_h1(lambda u: u.startswith("BAB I"))

    print("markers", abstrak_i, abstract_i, motto_i, persembahan_i, bab1_i)

    abstrak_blocks = []
    if abstrak_i is not None and abstract_i is not None:
        for b in raw[abstrak_i:abstract_i]:
            t = tx(b)
            if b.get("type") == "p" and (
                t.upper().startswith("DECKA")
                or t.upper().startswith("NIM ")
                or t == TITLE
                or t.upper().startswith("ATMOSPHERIC")
            ):
                continue
            abstrak_blocks.append(b)

    motto_blocks = []
    if motto_i is not None and persembahan_i is not None:
        motto_blocks = raw[motto_i:persembahan_i]
    elif motto_i is not None and bab1_i is not None:
        motto_blocks = raw[motto_i:bab1_i]

    body = raw[bab1_i:] if bab1_i is not None else raw

    # merge BAB + subtitle
    merged = [{"type": "title", "text": TITLE}]
    seq = [*abstrak_blocks, *motto_blocks, *body]
    i = 0
    while i < len(seq):
        b = seq[i]
        t = tx(b)
        if b.get("type") == "h1" and re.fullmatch(r"BAB [IVXLC]+", t, re.I) and i + 1 < len(seq):
            n = seq[i + 1]
            if n.get("type") in ("h1", "h1sub") and not re.match(r"^BAB\s", tx(n), re.I):
                merged.append(
                    {
                        "type": "h1",
                        "text": f"{t} {tx(n)}".strip(),
                        "style": b.get("style"),
                    }
                )
                i += 2
                continue
        if b.get("type") == "h1sub":
            # orphan subtitle
            merged.append({"type": "h1", "text": t, "style": b.get("style")})
            i += 1
            continue
        # promote Simpulan
        if b.get("type") == "p" and t in ("Simpulan", "Kesimpulan"):
            merged.append({"type": "h2", "text": "Simpulan"})
            i += 1
            continue
        if b.get("type") == "p" and t == "Saran":
            merged.append({"type": "h2", "text": "Saran"})
            i += 1
            continue
        # caption normalize later
        if b.get("type") == "h1" and t.upper().startswith("LAMPIRAN") and t.upper() != "LAMPIRAN":
            b = {**b, "type": "h2"}
        if b.get("type") == "h1" and t.startswith("Kesimpulan yang diperoleh"):
            b = {**b, "type": "p"}
        merged.append(b)
        i += 1

    # caption renumber + cleanup
    g = tnum = 0
    final = []
    for b in merged:
        b = dict(b)
        if b.get("type") == "caption" or (
            b.get("type") == "p" and re.match(r"^(Gambar|Figure|Tabel|Table)\b", tx(b), re.I)
        ):
            text = re.sub(r"\s+", " ", tx(b)).strip()
            m = re.match(r"^(Gambar|Figure|Tabel|Table)\s*(?:(\d+)\.)?\s*(.*)$", text, re.I)
            if m:
                kind, rest = m.group(1), (m.group(3) or "").strip(" .")
                rest = re.sub(r"^\d+\.\s*", "", rest)
                if kind.lower().startswith("gambar") or kind.lower().startswith("figure"):
                    g += 1
                    label, n = "Gambar", g
                else:
                    tnum += 1
                    label, n = "Tabel", tnum
                body = re.sub(r"\s+", " ", rest).strip()
                if body.endswith(")."):
                    body = body[:-1]
                if body and not body.endswith(".") and not body.endswith(")"):
                    body += "."
                final.append({"type": "caption", "text": f"{label} {n}. {body}".strip()})
            else:
                final.append({"type": "caption", "text": text})
            continue
        if b.get("type") == "h2":
            ht = re.sub(r"\s+", " ", tx(b)).strip()
            m = re.match(r"^Lampiran\s*([0-9]*)\s*\.?\s*(.*)$", ht, re.I)
            if m:
                num, rest = m.group(1), m.group(2).strip(" .")
                ht = f"Lampiran {num}. {rest}".strip() if num else (f"Lampiran. {rest}".strip() if rest else "Lampiran")
            ht = re.sub(r"\(\s+", "(", ht)
            ht = re.sub(r"\s+\)", ")", ht)
            b["text"] = ht
        if b.get("type") == "h1":
            b["text"] = re.sub(r"\s+", " ", tx(b)).strip()
            if re.fullmatch(r"BAB II", b["text"], re.I):
                b["text"] = "BAB II KAJIAN PUSTAKA"
        if b.get("type") == "li":
            b["text"] = re.sub(r"\s+", " ", tx(b)).strip()
        if b.get("type") == "p":
            b["text"] = re.sub(r"\s+", " ", tx(b)).strip()
        final.append(b)

    meta = {
        "title": TITLE,
        "author": "Decka Fadhila Tirta",
        "nim": "22306144046",
        "source": SRC.name,
        "blockCount": len(final),
        "imageCount": sum(1 for b in final if b.get("type") == "image"),
        "uniqueImages": len(saved),
        "textChars": sum(len(b.get("text", "")) for b in final),
        "listItems": sum(1 for b in final if b.get("type") == "li"),
        "filtered": True,
    }
    OUT_JSON.write_text(json.dumps({"meta": meta, "blocks": final}, ensure_ascii=False), encoding="utf-8")
    print(json.dumps(meta, indent=2, ensure_ascii=False))
    print("sample hierarchy:")
    for i, b in enumerate(final[:35]):
        print(i, b["type"], (b.get("text") or b.get("name", ""))[:90], "depth" if b.get("type") == "li" else "", b.get("depth", ""))
    print("li depths", {d: sum(1 for b in final if b.get("type")=="li" and b.get("depth")==d) for d in range(0,4)})


if __name__ == "__main__":
    main()
